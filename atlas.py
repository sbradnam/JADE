# -*- coding: utf-8 -*-
"""
Created on Wed Jan 15 14:30:18 2020

@author: Davide Laghi
"""
import docx
import win32com.client
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import pandas as pd


class Atlas():
    def __init__(self, template, name):
        """
        Atlas of plots for post-processing

        Parameters
        ----------
        template : Path/str
            Word template for atlas
        # lib : list/str
        #     libraries to post-process

        Returns
        -------
        None.

        """
        # if type(lib) == list and len(lib) > 1:
        #     self.comparison = True
        #     name = ''
        #     for l in lib:
        #         name = name+l+'-'
        #     del[name[-1]]
        # else:
        #     self.comparison = False
        #     name = lib

        # self.lib = str(lib)  # Libraries to post process
        self.name = name  # Name of the Atlas (from libraries)
        # Open The Atlas template
        doc = docx.Document(template)
        doc.add_heading('JADE ATLAS: '+name, level=0)
        self.outname = 'atlas_'+name  # Name for the outfile
        self.doc = doc  # Word Document

    def insert_img(self, img, width=Inches(7.5)):
        self.doc.add_picture(img, width=width)
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def build(self, images_path, libmanager, mat_settings):
        """
        TO BE USED FOR SPHERE LEAKAGE BENCHMARK

        Generate the atlas filling the world template with plots and headings

        Parameters
        ----------
        images_path : str/path
            Path to temporary folder containig the images

        lib_manager : libmanager.LibManager
            Library manager for conversions and name recovery.
        
        mat_settings : pd.DataFrame
            contains settings for Materials

        Returns
        -------
        None.

        """
        # Build Atlas
        images = []
        tallies = []
        for img in os.listdir(images_path):
            img_path = os.path.join(images_path, img)
            pieces = img.split('-')
            zaid = pieces[0]
            tally = pieces[-1].split('.')[0]
            zaidnum = zaid.split('.')[0]
            if tally not in tallies:
                tallies.append(tally)

            images.append({'tally': tally, 'zaid': zaid, 'img': img_path,
                           'num': zaidnum})

        images = pd.DataFrame(images)
        # Reorder atlas
        images['num'] = pd.to_numeric(images['num'].values, errors='coerce')
        images.set_index('tally', inplace=True)
        for tally in tallies:
            self.doc.add_heading('Tally N.'+str(tally), level=1)
            # Be sure of the reordering
            df = images.loc[tally].sort_values('num')
            for idx, row in df.iterrows():
                title = 'Zaid: '+row['zaid']
                try:
                    name, formula = libmanager.get_zaidname(row['zaid'])
                    title = title+' ('+name+' '+formula+')'
                except ValueError:  # A material is passed instead of zaid
                    matname = mat_settings.loc[row['zaid'], 'Name']
                    title = title+' ('+matname+')'

                self.doc.add_heading(title, level=2)
                self.insert_img(row['img'])

    def save(self, outpath, pdfprint=True):
        """
        Save word atlas and possibly export PDF

        Parameters
        ----------
        outpath : str/path
            path to export the atlas
        pdfprint : Boolean, optional
            If True export also in PDF format

        Returns
        -------
        None.

        """
        outpath_word = os.path.join(outpath, self.outname+'.docx')
        outpath_pdf = os.path.join(outpath, self.outname+'.pdf')

        self.doc.save(outpath_word)

        if pdfprint:
            in_file = outpath_word
            out_file = outpath_pdf

            word = win32com.client.Dispatch('Word.Application')
            doc = word.Documents.Open(in_file)
            doc.ExportAsFixedFormat(OutputFileName=out_file,
                                    # 17 = PDF output, 18=XPS output
                                    ExportFormat=17,
                                    OpenAfterExport=False,
                                    # 0=Print (higher res), 1=Screen (lower res)
                                    OptimizeFor=0,
   # 0=No bookmarks, 1=Heading bookmarks only, 2=bookmarks match word bookmarks
                                    CreateBookmarks=1,
                                    DocStructureTags=True)

            doc.Close()
            word.Quit()
